# -*- coding=utf-8 -*-
# @FileName  :word.py
# @Time      :2020/10/20 17:45
# @Author    :zyt

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


class Word(object):
    def __init__(self):
        self.document = Document()

    # 设置页面大小，页边距
    def set_section(self, width, height, top, bottom, left, right):
        section = self.document.sections[0]
        section.page_width = Cm(width)
        section.page_height = Cm(height)
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

    # 添加标题
    def add_title_text(self, space_before, space_after, line_space, position):
        title = self.document.add_paragraph()
        title_format = title.paragraph_format

        if position == 0:
            title_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 标题居左
        elif position == 1:
            title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题居中
        elif position == 2:
            title_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif position == 3:
            title_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        title_format.space_before = Pt(space_before)
        title_format.space_after = Pt(space_after)
        title_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        title_format.line_spacing = Pt(line_space)
        return title

    # 添加段落并设置格式()
    def add_para_text(self, space_before, space_after, line_space):
        paragraph = self.document.add_paragraph()  # 段落
        par_format = paragraph.paragraph_format
        par_format.first_line_indent = Pt(20)  # 首行缩进
        par_format.space_before = Pt(space_before)
        par_format.space_after = Pt(space_after)
        par_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        par_format.line_spacing = Pt(line_space)
        return paragraph

    @staticmethod
    def add_run_text(paragraph, text, size, bold, underline, font_name):
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.bold = bold
        run.font.underline = underline
        run.font.color.rgb = RGBColor(0, 0, 0)
